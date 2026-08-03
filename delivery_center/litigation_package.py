"""根据案件信息与法律分析生成五份诉讼交付 DOCX。"""

from __future__ import annotations

import json
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ai.llm_client import LLMMessage, ProviderLLMClient
from lawsuit_generator.civil_complaint import CivilComplaintGenerator


class LitigationPackageError(RuntimeError):
    """诉讼材料包数据不完整或生成失败。"""


class LitigationPackageGenerator:
    REQUIRED_FIELDS = ("evidence_catalog", "evidence_explanations", "legal_basis_list", "litigation_risk_analysis")

    def __init__(self, *, llm_client=None, complaint_generator=None) -> None:
        self._llm_client = llm_client or ProviderLLMClient("deepseek")
        self._complaint_generator = complaint_generator or CivilComplaintGenerator()

    def generate(self, case, analysis: dict, records: list) -> dict[str, bytes]:
        context = {
            "case": {"name": case.name, "parties": case.parties, "case_type": case.case_type},
            "legal_analysis": analysis,
            "case_records": [{"type": record.record_type, "title": record.title, "content": record.content} for record in records],
        }
        prompt = f"""你是中国律师诉讼材料整理助手。请基于提供的案件资料输出严格 JSON，不得虚构证据或法条。
顶层字段必须为 evidence_catalog、evidence_explanations、legal_basis_list、litigation_risk_analysis，均为数组。
evidence_catalog 每项应包含 number、name、source、purpose、status；缺失材料应标记“待补充”。
所有不确定信息必须明确标注待核实，材料须经律师审核。

案件资料：{json.dumps(context, ensure_ascii=False)}
"""
        response = self._llm_client.complete([LLMMessage("user", prompt)], response_format="json")
        try:
            package_data = json.loads(response.content)
        except (TypeError, json.JSONDecodeError) as exc:
            raise LitigationPackageError("DeepSeek 未返回有效的诉讼材料包结构。") from exc
        missing = [field for field in self.REQUIRED_FIELDS if field not in package_data]
        if missing:
            raise LitigationPackageError(f"诉讼材料包缺少字段：{', '.join(missing)}")

        complaint_context = json.dumps(context, ensure_ascii=False)
        return {
            "民事起诉状.docx": self._complaint_generator.generate(complaint_context),
            "证据目录.docx": self._build_evidence_catalog(case, package_data["evidence_catalog"]),
            "证据说明.docx": self._build_sections(case, "证据说明", package_data["evidence_explanations"]),
            "法律依据清单.docx": self._build_sections(case, "法律依据清单", package_data["legal_basis_list"]),
            "诉讼风险分析.docx": self._build_sections(case, "诉讼风险分析", package_data["litigation_risk_analysis"]),
        }

    @staticmethod
    def _new_document(case, title: str) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        document.add_paragraph(f"案件：{case.name}\n当事人：{case.parties}\n案件类型：{case.case_type}")
        return document

    @staticmethod
    def _save(document: Document) -> bytes:
        document.add_paragraph("律师审核提示：本材料由 AI 辅助整理，提交或对外使用前须核对事实、证据原件、法条效力及程序要求。")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @classmethod
    def _build_evidence_catalog(cls, case, items: list) -> bytes:
        document = cls._new_document(case, "证据目录")
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.autofit = False
        headers = ("序号", "证据名称", "来源", "证明目的", "状态")
        for cell, value in zip(table.rows[0].cells, headers):
            cell.text = value
            cell.paragraphs[0].runs[0].bold = True
        for index, item in enumerate(items, start=1):
            row = table.add_row().cells
            values = (
                item.get("number", index) if isinstance(item, dict) else index,
                item.get("name", "待补充") if isinstance(item, dict) else cls._stringify(item),
                item.get("source", "待补充") if isinstance(item, dict) else "待补充",
                item.get("purpose", "待核实") if isinstance(item, dict) else "待核实",
                item.get("status", "待补充") if isinstance(item, dict) else "待补充",
            )
            for cell, value in zip(row, values):
                cell.text = str(value)
        return cls._save(document)

    @classmethod
    def _build_sections(cls, case, title: str, items: list) -> bytes:
        document = cls._new_document(case, title)
        if not items:
            document.add_paragraph("暂无可确认内容，待律师补充。")
        for index, item in enumerate(items, start=1):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(cls._stringify(item))
        return cls._save(document)

    @staticmethod
    def _stringify(value) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return "；".join(f"{key}：{LitigationPackageGenerator._stringify(item)}" for key, item in value.items())
        if isinstance(value, list):
            return "；".join(LitigationPackageGenerator._stringify(item) for item in value)
        return "" if value is None else str(value)
