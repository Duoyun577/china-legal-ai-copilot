"""生成合同原条款与 AI 修订条款的修改说明 DOCX。"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class ContractDiffGenerator:
    """根据原合同、修订版 DOCX 和现有风险结果生成修改说明。"""

    def generate(self, original_text: str, revised_contract: bytes, review) -> bytes:
        revised_document = Document(BytesIO(revised_contract))
        revised_clauses = [paragraph.text for paragraph in revised_document.paragraphs if "【AI修改】" in paragraph.text]
        document = self._new_document(review.contract_path.name)
        if not review.risks:
            document.add_paragraph("现有规则未发现需要说明的合同风险修改。")
        for index, risk in enumerate(review.risks, start=1):
            document.add_heading(f"{index}. {risk.name}（{risk.rule_id}）", level=1)
            original_clause = self._find_original_clause(original_text, risk.matched_keywords)
            revised_clause = revised_clauses[index - 1] if index <= len(revised_clauses) else "未能自动对应修改条款，需律师人工复核。"
            self._add_field(document, "原条款", original_clause)
            self._add_field(document, "修改后条款", revised_clause)
            self._add_field(document, "修改理由", risk.legal_issue)
            self._add_field(document, "降低风险", risk.suggestion)
        document.add_paragraph("提示：本说明基于规则命中与 AI 修改标记自动生成，条款对应关系及修改效果须由律师复核。")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _new_document(contract_name: str) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        heading_style = document.styles["Heading 1"]
        heading_style.font.color.rgb = RGBColor(46, 116, 181)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("合同修改说明")
        run.bold = True
        run.font.size = Pt(22)
        subtitle = document.add_paragraph(f"合同：{contract_name}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document

    @staticmethod
    def _find_original_clause(text: str, keywords: list[str]) -> str:
        blocks = [block.strip() for block in text.splitlines() if block.strip()]
        matches = [block for block in blocks if any(keyword in block for keyword in keywords if keyword)]
        return "\n".join(matches[:3]) or "未能根据风险关键词定位原条款，需律师人工复核。"

    @staticmethod
    def _add_field(document: Document, label: str, value: str) -> None:
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        paragraph.add_run(value)

