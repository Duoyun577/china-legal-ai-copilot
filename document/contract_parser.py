"""将上传的合同文档解析为 ContractReviewService 可读取的纯文本。"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocumentParseError(ValueError):
    """上传文档无法转换为可审查文本。"""


@dataclass(frozen=True)
class ParsedContract:
    """统一的合同文本解析结果。"""

    filename: str
    file_type: str
    text: str


class ContractDocumentParser:
    """根据文件扩展名解析 TXT、DOCX 和文本型 PDF。"""

    supported_extensions = {".txt", ".docx", ".pdf"}

    def parse(self, filename: str, content: bytes) -> ParsedContract:
        safe_name = Path(filename).name or "uploaded_contract.txt"
        extension = Path(safe_name).suffix.lower()
        if extension not in self.supported_extensions:
            raise DocumentParseError("仅支持 TXT、DOCX 和 PDF 合同。")
        if not content:
            raise DocumentParseError("上传的合同文件为空。")

        parsers = {
            ".txt": self._parse_txt,
            ".docx": self._parse_docx,
            ".pdf": self._parse_pdf,
        }
        try:
            text = parsers[extension](content)
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"{extension[1:].upper()} 文件解析失败。") from exc

        normalized_text = text.strip()
        if not normalized_text:
            detail = "PDF 可能是扫描件，当前仅支持包含文本层的 PDF。" if extension == ".pdf" else "文档中没有可审查的文本。"
            raise DocumentParseError(detail)
        return ParsedContract(filename=safe_name, file_type=extension[1:], text=normalized_text)

    @staticmethod
    def _parse_txt(content: bytes) -> str:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("TXT 文件必须使用 UTF-8 编码。") from exc

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        document = Document(BytesIO(content))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
        return "\n".join(blocks)

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(page for page in pages if page)
