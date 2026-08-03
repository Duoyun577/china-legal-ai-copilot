"""使用现有 AI Provider 生成带标记的完整修订版合同 DOCX。"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ai.llm_client import LLMMessage, ProviderLLMClient


class ContractRewriteError(RuntimeError):
    """AI 未能返回可交付的完整修订合同。"""


class AIContractRewriter:
    """根据原合同和现有风险结果调用 DeepSeek 并输出完整 DOCX。"""

    def __init__(self, llm_client=None) -> None:
        self._llm_client = llm_client or ProviderLLMClient("deepseek")

    def rewrite(self, contract_text: str, review) -> bytes:
        prompt = self._build_prompt(contract_text, review)
        response = self._llm_client.complete([LLMMessage("user", prompt)], response_format="text")
        revised_text = self._clean_response(response.content)
        self._validate(revised_text, contract_text, review)
        return self._build_docx(revised_text)

    @staticmethod
    def _build_prompt(contract_text: str, review) -> str:
        risks = "\n".join(
            f"- {risk.rule_id}｜{risk.name}｜{risk.risk_level}｜问题：{risk.legal_issue}｜建议：{risk.suggestion}"
            for risk in review.risks
        )
        return f"""你是中国律师合同修订助手。请输出一份完整的修订版合同纯文本。

强制要求：
1. 严格保留原合同的标题、章节顺序和无风险条款，不得摘要或省略。
2. 只修改下列风险直接涉及的条款，不得擅自改变其他商业条件。
3. 每个被修改的条款开头必须添加【AI修改】；未修改条款不得添加该标记。
4. 输出完整合同正文，不要输出解释、前言、Markdown 代码围栏或修改清单。
5. 不得编造当事人信息、金额、日期或法律依据；缺失内容使用明确的待协商占位说明。

风险结果：
{risks or '现有规则未发现风险，不得修改合同。'}

原合同全文：
{contract_text}
"""

    @staticmethod
    def _clean_response(content: str) -> str:
        text = content.strip()
        fence = re.fullmatch(r"```(?:text|markdown)?\s*(.*?)\s*```", text, flags=re.DOTALL | re.IGNORECASE)
        return (fence.group(1) if fence else text).strip()

    @staticmethod
    def _validate(text: str, original_text: str, review) -> None:
        if review.risks and "【AI修改】" not in text:
            raise ContractRewriteError("AI 修订合同未标记修改位置。")
        minimum_length = max(1, int(len(original_text.strip()) * 0.65))
        if not text or len(text) < minimum_length:
            raise ContractRewriteError("AI 返回的修订合同为空或明显不完整。")

    @staticmethod
    def _build_docx(text: str) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        for index, block in enumerate(text.splitlines()):
            block = block.strip()
            if not block:
                continue
            paragraph = document.add_paragraph()
            if index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if block.startswith("【AI修改】"):
                marker = paragraph.add_run("【AI修改】")
                marker.bold = True
                marker.font.color.rgb = RGBColor(192, 57, 43)
                paragraph.add_run(block.removeprefix("【AI修改】").lstrip())
            else:
                run = paragraph.add_run(block)
                if index == 0:
                    run.bold = True
                    run.font.size = Pt(18)
        output = BytesIO()
        document.save(output)
        return output.getvalue()
