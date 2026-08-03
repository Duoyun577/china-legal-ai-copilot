"""生成《合同风险审查及修改建议书》DOCX。"""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from legal_assistant.citation_utils import format_citation, risk_citations


class ContractReviewReportGenerator:
    """将现有 ContractReviewService 结果转换为正式 DOCX 建议书。"""

    def generate(self, review, contract_text: str) -> bytes:
        document = Document()
        self._configure(document)
        self._add_header(document)
        self._add_basic_info(document, review)
        self._add_overall_risk(document, review)
        self._add_risk_list(document, review, contract_text)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _configure(document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
            style = document.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(46, 116, 181)

    @staticmethod
    def _add_header(document: Document) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run("合同风险审查及修改建议书")
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(22)
        subtitle = document.add_paragraph("China Legal AI Copilot · 待律师复核")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(18)

    def _add_basic_info(self, document: Document, review) -> None:
        document.add_heading("一、合同基本信息", level=1)
        rows = [
            ("合同名称", review.contract_path.name),
            ("合同类型", review.contract_type),
            ("审查日期", date.today().isoformat()),
            ("审查方式", "规则引擎与本地法律依据辅助审查"),
        ]
        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.autofit = False
        for row, (label, value) in zip(table.rows, rows):
            row.cells[0].width = Inches(1.18)
            row.cells[1].width = Inches(5.32)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].text = label
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].bold = True
        self._set_table_geometry(table, (1700, 7660))

    def _add_overall_risk(self, document: Document, review) -> None:
        document.add_heading("二、总体风险等级", level=1)
        level = self._overall_level(review)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{level} · 共发现 {len(review.risks)} 项风险")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(192, 57, 43) if level == "HIGH" else RGBColor(191, 130, 0)

    def _add_risk_list(self, document: Document, review, contract_text: str) -> None:
        document.add_heading("三、风险列表与修改建议", level=1)
        if not review.risks:
            document.add_paragraph("现有规则未发现风险，仍建议由律师结合交易背景复核。")
            return
        for index, risk in enumerate(review.risks, start=1):
            document.add_heading(f"{index}. {risk.name}（{risk.rule_id} / {risk.risk_level}）", level=2)
            self._add_labeled_paragraph(document, "原条款", self._find_original_clause(contract_text, risk.matched_keywords))
            self._add_labeled_paragraph(document, "风险说明", f"{risk.description} {risk.legal_issue}")
            bases = review.legal_basis_by_rule.get(risk.rule_id, [])
            legal_text = "\n\n".join(format_citation(item) for item in risk_citations(risk, bases))
            self._add_labeled_paragraph(document, "法律依据", legal_text or "暂无本地法律库关联条目，需律师进一步检索核验。")
            self._add_labeled_paragraph(document, "修改建议", risk.suggestion)

    @staticmethod
    def _add_labeled_paragraph(document: Document, label: str, value: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        paragraph.add_run(value)

    @staticmethod
    def _find_original_clause(text: str, keywords: list[str]) -> str:
        blocks = [block.strip() for block in text.splitlines() if block.strip()]
        matches = [block for block in blocks if any(keyword in block for keyword in keywords if keyword)]
        return "\n".join(matches[:3]) or "未能根据规则关键词定位具体条款，需人工复核全文。"

    @staticmethod
    def _overall_level(review) -> str:
        if any(risk.risk_level == "HIGH" for risk in review.risks):
            return "HIGH"
        if any(risk.risk_level == "MIDDLE" for risk in review.risks):
            return "MIDDLE"
        return "LOW"

    @staticmethod
    def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
        table_element = table._tbl
        properties = table_element.tblPr
        width = properties.first_child_found_in("w:tblW")
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(sum(widths)))
        grid = table_element.tblGrid
        for child in list(grid):
            grid.remove(child)
        for value in widths:
            column = OxmlElement("w:gridCol")
            column.set(qn("w:w"), str(value))
            grid.append(column)
        for row in table.rows:
            for cell, value in zip(row.cells, widths):
                cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                cell_width.set(qn("w:type"), "dxa")
                cell_width.set(qn("w:w"), str(value))
