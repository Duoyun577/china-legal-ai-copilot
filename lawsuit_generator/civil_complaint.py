"""使用现有 lawsuit_drafting Prompt 生成民事起诉状 DOCX。"""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ai.llm_client import LLMMessage, ProviderLLMClient
from legal_assistant.citation_utils import citation_dict, format_citation
from legal_assistant.legal_search_adapter import LegalKnowledgeSearch


class CivilComplaintError(RuntimeError):
    """民事起诉状无法生成。"""


class CivilComplaintGenerator:
    def __init__(self, llm_client=None, prompt_path: Path | None = None, *, legal_search=None) -> None:
        root = Path(__file__).resolve().parents[1]
        self._llm_client = llm_client or ProviderLLMClient("deepseek")
        self._prompt_path = prompt_path or root / "ai" / "prompts" / "lawsuit_drafting.md"
        self._legal_search = legal_search or LegalKnowledgeSearch()

    def generate(self, case_facts: str) -> bytes:
        facts = case_facts.strip()
        if not facts:
            raise CivilComplaintError("请输入案件事实。")
        system_prompt = self._prompt_path.read_text(encoding="utf-8")
        legal_basis = [citation_dict(item) for item in self._legal_search.search(facts, top_k=5)]
        if not legal_basis:
            raise CivilComplaintError("本地法律知识库未检索到可核验法条，无法生成起诉状。")
        payload = {
            "case_type": "民事纠纷", "parties": [], "facts": [facts], "claims": [],
            "evidence": [], "legal_basis": legal_basis, "court_information": {},
        }
        prompt = f"{system_prompt}\n\n请根据以下输入生成符合上述结构的 JSON。legal_basis 仅可使用输入法条，不得编造；另输出 related_legal_basis_suggestions 数组：\n{json.dumps(payload, ensure_ascii=False)}"
        response = self._llm_client.complete([LLMMessage("user", prompt)], response_format="json")
        try:
            result = json.loads(response.content)
        except (TypeError, json.JSONDecodeError) as exc:
            raise CivilComplaintError("DeepSeek 未返回有效的起诉状结构。") from exc
        if not isinstance(result, dict) or not result.get("document_type"):
            raise CivilComplaintError("起诉状结构不完整，请补充案件信息后重试。")
        result["related_legal_basis_suggestions"] = legal_basis
        return self._build_docx(result)

    @staticmethod
    def _build_docx(result: dict) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "SimSun"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("民事起诉状")
        run.bold = True
        run.font.size = Pt(22)
        CivilComplaintGenerator._add_section(document, "当事人信息", result.get("parties", []))
        CivilComplaintGenerator._add_section(document, "诉讼请求", result.get("claims", []), numbered=True)
        CivilComplaintGenerator._add_section(document, "事实与理由", result.get("facts_and_reasons", []))
        CivilComplaintGenerator._add_section(document, "证据清单", result.get("evidence_list", []), numbered=True)
        CivilComplaintGenerator._add_section(document, "法律依据", result.get("legal_basis", []))
        CivilComplaintGenerator._add_section(
            document,
            "相关法律依据建议",
            [format_citation(item) for item in result.get("related_legal_basis_suggestions", [])],
        )
        CivilComplaintGenerator._add_section(document, "待核实事项", result.get("procedural_uncertainties", []))

        court = CivilComplaintGenerator._stringify(result.get("court")) or "【待确认有管辖权的人民法院】"
        document.add_paragraph(f"此致\n{court}")
        signature = document.add_paragraph("具状人：【待填写】\n日期：【待填写】")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_paragraph("提示：本起诉状为 AI 辅助初稿，须经执业律师审核并由当事人确认后使用。")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _add_section(document: Document, heading: str, values, *, numbered: bool = False) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        items = values if isinstance(values, list) else [values]
        if not items:
            document.add_paragraph("【待补充】")
            return
        for index, item in enumerate(items, start=1):
            prefix = f"{index}. " if numbered else ""
            document.add_paragraph(prefix + CivilComplaintGenerator._stringify(item))

    @staticmethod
    def _stringify(value) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return "；".join(f"{key}：{CivilComplaintGenerator._stringify(item)}" for key, item in value.items())
        if isinstance(value, list):
            return "；".join(CivilComplaintGenerator._stringify(item) for item in value)
        return "" if value is None else str(value)
