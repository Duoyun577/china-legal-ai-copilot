"""面向律师工作和法院提交的双版本民事诉状生成服务。"""

from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ai.llm_client import LLMMessage, ProviderLLMClient
from legal_assistant.legal_reference_validator import LegalReferenceValidator


class PleadingGenerationError(RuntimeError):
    """诉讼方案或诉状无法生成。"""


@dataclass(frozen=True)
class PleadingDocuments:
    lawyer_version: bytes
    court_version: bytes
    pleading_plan: dict


class LitigationPleadingService:
    REQUIRED_FIELDS = (
        "cause_of_action", "dispute_issues", "claims", "facts_and_reasons", "parties",
        "court", "evidence_system", "risks", "litigation_strategy", "procedural_uncertainties",
    )
    FORBIDDEN_COURT_MARKERS = ("AI", "人工智能", "模型生成", "辅助初稿", "律师工作版")

    def __init__(self, llm_client=None) -> None:
        self._llm_client = llm_client or ProviderLLMClient("deepseek")

    def generate(self, context: dict) -> PleadingDocuments:
        legal_basis = list(context.get("verified_legal_basis", []))
        if not legal_basis:
            raise PleadingGenerationError("诉状缺少可核验法律依据。")
        try:
            validations = LegalReferenceValidator().validate_many(legal_basis, strict=True)
        except ValueError as exc:
            raise PleadingGenerationError(str(exc)) from exc
        context = {**context, "legal_citation_validation": validations}
        prompt = f"""你是中国执业律师的民事诉讼文书助手。请仅输出严格 JSON，并按以下顺序形成可复核诉讼方案：
案件记忆、案件分析、案由判断、争议焦点、诉讼请求、法律依据、类案参考、证据体系、生成诉状。
顶层字段必须包含 cause_of_action（字符串）、dispute_issues（数组）、claims（数组）、facts_and_reasons（数组）、parties（数组）、court（字符串）、evidence_system（数组）、risks（数组）、litigation_strategy（数组）、procedural_uncertainties（数组）。
evidence_system 每项包含 name、source、purpose、status。只能使用输入中的 verified_legal_basis 和 verified_similar_cases，不得编造事实、证据、法条、法院或当事人身份；未知内容使用【待核实】或【待填写】。诉讼请求应具体、可执行，金额及计算方式不明时明确留待核实。
案件完整资料：{json.dumps(context, ensure_ascii=False)}
"""
        response = self._llm_client.complete([LLMMessage("user", prompt)], response_format="json")
        try:
            plan = json.loads(response.content)
        except (TypeError, json.JSONDecodeError) as exc:
            raise PleadingGenerationError("模型未返回有效的诉讼方案。") from exc
        missing = [field for field in self.REQUIRED_FIELDS if field not in plan]
        if missing:
            raise PleadingGenerationError(f"诉讼方案缺少字段：{', '.join(missing)}")
        plan["legal_basis"] = legal_basis
        plan["legal_citation_validation"] = validations
        plan["similar_case_references"] = list(context.get("verified_similar_cases", []))
        lawyer = self._build_lawyer_version(context, plan)
        court = self._build_court_version(plan)
        self._validate_court_version(court)
        return PleadingDocuments(lawyer, court, plan)

    @classmethod
    def _new_document(cls, title: str, *, formal: bool = False) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.15 if formal else 1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "SimSun"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        return document

    @classmethod
    def _build_lawyer_version(cls, context: dict, plan: dict) -> bytes:
        document = cls._new_document("民事诉讼方案及起诉状（律师工作版）")
        cls._section(document, "一、案件记忆", context.get("案件长期记忆", {}))
        cls._section(document, "二、案件分析", context.get("案件法律分析", {}))
        cls._section(document, "三、案由判断", plan["cause_of_action"])
        cls._section(document, "四、争议焦点", plan["dispute_issues"], numbered=True)
        cls._section(document, "五、诉讼请求", plan["claims"], numbered=True)
        cls._section(document, "六、法律依据", plan["legal_basis"], numbered=True)
        cls._section(document, "七、类案参考", plan["similar_case_references"], numbered=True)
        cls._section(document, "八、证据体系", plan["evidence_system"], numbered=True)
        cls._section(document, "九、诉讼风险", plan["risks"], numbered=True)
        cls._section(document, "十、诉讼策略", plan["litigation_strategy"], numbered=True)
        cls._section(document, "十一、待核实事项", plan["procedural_uncertainties"], numbered=True)
        document.add_paragraph("内部工作提示：提交前应由承办律师核验主体、管辖、请求金额、证据原件及法条时效性。")
        return cls._save(document)

    @classmethod
    def _build_court_version(cls, plan: dict) -> bytes:
        document = cls._new_document("民事起诉状", formal=True)
        for party in plan["parties"] or ["原告：【待填写】", "被告：【待填写】"]:
            document.add_paragraph(cls._stringify(party))
        cls._section(document, "诉讼请求", plan["claims"], numbered=True)
        cls._section(document, "事实与理由", plan["facts_and_reasons"])
        references = "、".join(cls._legal_reference(item) for item in plan["legal_basis"] if cls._legal_reference(item))
        if references:
            document.add_paragraph(f"综上，依据{references}及相关规定，请求人民法院依法支持原告的诉讼请求。")
        court = cls._stringify(plan["court"]) or "【待确认有管辖权的人民法院】"
        document.add_paragraph(f"此致\n{court}")
        signature = document.add_paragraph("具状人：【待填写】\n日期：【待填写】")
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        evidence_names = [item.get("name", "【待补充】") if isinstance(item, dict) else cls._stringify(item) for item in plan["evidence_system"]]
        document.add_paragraph(f"附：\n1. 本起诉状副本【待填写】份；\n2. 证据材料目录：{cls._stringify(evidence_names) or '【待补充】'}。")
        return cls._save(document)

    @classmethod
    def _section(cls, document: Document, heading: str, values, *, numbered: bool = False) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        items = values if isinstance(values, list) else [values]
        if not items:
            items = ["【待补充】"]
        for index, item in enumerate(items, start=1):
            prefix = f"{index}. " if numbered else ""
            document.add_paragraph(prefix + cls._stringify(item))

    @staticmethod
    def _stringify(value) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return "；".join(f"{key}：{LitigationPleadingService._stringify(item)}" for key, item in value.items())
        if isinstance(value, list):
            return "；".join(LitigationPleadingService._stringify(item) for item in value)
        return "" if value is None else str(value)

    @staticmethod
    def _legal_reference(value) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return str(value.get("legal_basis") or f"《{value.get('law_name', '')}》{value.get('article', '')}").strip()
        return ""

    @staticmethod
    def _save(document: Document) -> bytes:
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @classmethod
    def _validate_court_version(cls, content: bytes) -> None:
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        marker = next((item for item in cls.FORBIDDEN_COURT_MARKERS if item in text), None)
        if marker:
            raise PleadingGenerationError(f"法院提交版包含禁止标识：{marker}")
