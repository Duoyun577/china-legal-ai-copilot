"""生成并渲染《案件法律分析报告》。"""

from __future__ import annotations

import json
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ai.llm_client import LLMMessage, ProviderLLMClient
from legal_assistant.citation_utils import citation_dict, extract_citations, format_citation
from legal_assistant.legal_search_adapter import LegalKnowledgeSearch


class CaseAnalysisError(RuntimeError):
    """案件法律分析无法生成。"""


class CaseLegalAnalysisReportGenerator:
    FIELDS = (
        "case_facts", "legal_relationships", "dispute_issues", "legal_basis",
        "risk_analysis", "litigation_strategy", "next_steps",
    )

    def __init__(self, llm_client=None, *, legal_search=None) -> None:
        self._llm_client = llm_client or ProviderLLMClient("deepseek")
        self._legal_search = legal_search or LegalKnowledgeSearch()

    def generate(self, case, consultation_records: list) -> tuple[dict, bytes]:
        if not consultation_records:
            raise CaseAnalysisError("当前案件尚无法律咨询记录，请先完成法律咨询。")
        context = {
            "case": {"name": case.name, "parties": case.parties, "case_type": case.case_type},
            "consultations": [{"title": record.title, "content": record.content} for record in consultation_records],
        }
        legal_basis = extract_citations(context)
        if not legal_basis:
            query = " ".join((case.name, case.parties, case.case_type, json.dumps(context, ensure_ascii=False)))
            legal_basis = [citation_dict(item) for item in self._legal_search.search(query, top_k=5)]
        if not legal_basis:
            raise CaseAnalysisError("本地法律知识库未检索到可核验法条，无法生成案件分析报告。")
        context["verified_legal_basis"] = legal_basis
        prompt = f"""你是中国律师案件分析辅助助手。请根据案件信息和已保存的咨询记录输出严格 JSON。
顶层字段必须为：case_facts、legal_relationships、dispute_issues、legal_basis、risk_analysis、litigation_strategy、next_steps。
各字段可为字符串或数组。legal_basis 只能使用 verified_legal_basis 中的法条，并保留法律名称、条文编号、条文内容和来源文件。不得虚构事实、证据或法条；信息不足须明确标为待核实；结论须经律师复核。

案件资料：{json.dumps(context, ensure_ascii=False)}
"""
        response = self._llm_client.complete([LLMMessage("user", prompt)], response_format="json")
        try:
            analysis = json.loads(response.content)
        except (TypeError, json.JSONDecodeError) as exc:
            raise CaseAnalysisError("DeepSeek 未返回有效的案件法律分析结构。") from exc
        missing = [field for field in self.FIELDS if field not in analysis]
        if missing:
            raise CaseAnalysisError(f"案件法律分析缺少字段：{', '.join(missing)}")
        analysis["legal_basis"] = legal_basis
        return analysis, self._build_docx(case, analysis)

    @staticmethod
    def _build_docx(case, analysis: dict) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        for style_name, size in (("Heading 1", 16), ("Heading 2", 13)):
            style = document.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(46, 116, 181)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("案件法律分析报告")
        title_run.bold = True
        title_run.font.size = Pt(22)
        metadata = document.add_paragraph(f"案件：{case.name}\n当事人：{case.parties}\n案件类型：{case.case_type}")
        metadata.paragraph_format.space_after = Pt(18)
        sections = (
            ("一、案件事实整理", "case_facts"),
            ("二、法律关系分析", "legal_relationships"),
            ("三、争议焦点", "dispute_issues"),
            ("四、法律依据", "legal_basis"),
            ("五、风险分析", "risk_analysis"),
            ("六、诉讼策略", "litigation_strategy"),
            ("七、下一步建议", "next_steps"),
        )
        for heading, field in sections:
            document.add_heading(heading, level=1)
            value = analysis[field]
            items = value if isinstance(value, list) else [value]
            for item in items:
                text = format_citation(item) if field == "legal_basis" and isinstance(item, dict) else CaseLegalAnalysisReportGenerator._stringify(item)
                document.add_paragraph(text)
        document.add_paragraph("提示：本报告为 AI 辅助分析，须由执业律师结合原始证据、时效及程序事项复核。")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _stringify(value) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            return "；".join(f"{key}：{CaseLegalAnalysisReportGenerator._stringify(item)}" for key, item in value.items())
        if isinstance(value, list):
            return "；".join(CaseLegalAnalysisReportGenerator._stringify(item) for item in value)
        return "" if value is None else str(value)
