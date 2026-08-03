from io import BytesIO
import json
from pathlib import Path

from docx import Document

from ai.llm_client import LLMResponse
from case_manager import CaseManager
from case_manager.workflow import CaseWorkflow
from delivery_center import LitigationPackageGenerator


PACKAGE_DATA = {
    "evidence_catalog": [
        {"number": 1, "name": "服务合同", "source": "委托人提供", "purpose": "证明合同关系", "status": "已有复印件"},
        {"number": 2, "name": "付款记录", "source": "银行", "purpose": "证明付款情况", "status": "待补充"},
    ],
    "evidence_explanations": ["服务合同用于证明双方权利义务。", "付款记录需进一步调取核实。"],
    "legal_basis_list": ["《中华人民共和国民法典》第五百七十七条。"],
    "litigation_risk_analysis": ["被告履行情况和损失金额仍需举证。"],
}


class FakeLLMClient:
    def complete(self, messages, *, response_format="text") -> LLMResponse:
        return LLMResponse(content=json.dumps(PACKAGE_DATA, ensure_ascii=False), model="test", is_mock=True)


class FakeComplaintGenerator:
    def generate(self, case_facts: str) -> bytes:
        document = Document()
        document.add_paragraph("民事起诉状")
        document.add_paragraph(case_facts[:100])
        output = BytesIO()
        document.save(output)
        return output.getvalue()


def docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def test_litigation_package_generates_five_complete_docx_files(tmp_path: Path) -> None:
    manager = CaseManager(tmp_path / "cases.db")
    case = manager.create_case("服务合同纠纷", "甲公司 / 乙公司", "合同纠纷")
    generator = LitigationPackageGenerator(llm_client=FakeLLMClient(), complaint_generator=FakeComplaintGenerator())

    documents = generator.generate(case, {"case_facts": ["被告逾期付款"]}, [])

    assert set(documents) == {"民事起诉状.docx", "证据目录.docx", "证据说明.docx", "法律依据清单.docx", "诉讼风险分析.docx"}
    assert all(content.startswith(b"PK") for content in documents.values())
    assert "证明目的" in docx_text(documents["证据目录.docx"])
    assert "待补充" in docx_text(documents["证据目录.docx"])
    assert "律师审核提示" in docx_text(documents["诉讼风险分析.docx"])


def test_delivery_package_is_saved_and_each_file_enters_timeline(tmp_path: Path) -> None:
    manager = CaseManager(tmp_path / "cases.db")
    case = manager.create_case("服务合同纠纷", "甲公司 / 乙公司", "合同纠纷")
    manager.add_record(case.case_id, "case_legal_analysis", "案件法律分析报告", {"case_facts": ["逾期付款"]})
    package_generator = LitigationPackageGenerator(llm_client=FakeLLMClient(), complaint_generator=FakeComplaintGenerator())
    workflow = CaseWorkflow(manager, package_generator=package_generator)

    documents = workflow.generate_litigation_package(case.case_id)

    assert len(documents) == 5
    assert len(manager.list_files(case.case_id)) == 5
    generated_events = [event for event in manager.list_events(case.case_id) if event.event_type == "generated_file"]
    assert len(generated_events) == 5
    assert {event.details["filename"] for event in generated_events} == set(documents)
    assert any(record.record_type == "delivery_package" for record in manager.list_records(case.case_id))
