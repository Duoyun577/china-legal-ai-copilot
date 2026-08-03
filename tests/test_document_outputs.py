from io import BytesIO
from pathlib import Path
import sys

from docx import Document
import pytest


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "app"))

from contract_review_service import ContractReviewService
from document.contract_rewriter import AIContractRewriter, ContractRewriteError
from document.contract_diff import ContractDiffGenerator
from document.report_generator import ContractReviewReportGenerator
from ai.llm_client import LLMResponse
from streamlit_app import (
    MAX_UPLOAD_BYTES,
    UploadValidationError,
    ai_revision_summary,
    contract_risk_score,
    generate_download_documents,
    risk_level_counts,
    validate_upload,
)


def document_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


class FakeLLMClient:
    def __init__(self, content: str) -> None:
        self.content = content
        self.messages = []

    def complete(self, messages, *, response_format="text") -> LLMResponse:
        self.messages = messages
        return LLMResponse(content=self.content, model="test-provider", is_mock=True)


def test_review_report_docx_contains_required_sections() -> None:
    path = Path("evaluation/test_contracts/software_service_contract.txt")
    contract_text = path.read_text(encoding="utf-8")
    review = ContractReviewService().review(path)

    result = ContractReviewReportGenerator().generate(review, contract_text)
    text = document_text(result)

    assert text.startswith("合同风险审查及修改建议书")
    assert "合同基本信息" in text
    assert "总体风险等级" in text
    assert "风险列表与修改建议" in text
    assert "原条款：" in text
    assert "风险说明：" in text
    assert "法律依据：" in text
    assert "法律名称：" in text
    assert "条文编号：" in text
    assert "条文内容：" in text
    assert "来源文件：knowledge_base/laws/civil_code_contract.json" in text
    assert text.count("来源文件：") >= len(review.risks)
    assert "修改建议：" in text
    assert "CR-009" in text


def test_ai_rewriter_outputs_complete_marked_contract() -> None:
    original = "软件服务合同\n第一条 服务内容保持不变。\n第二条 付款期限由双方另行协商。"
    revised = "软件服务合同\n第一条 服务内容保持不变。\n【AI修改】第二条 付款期限为验收合格并收到发票后十个工作日内。"
    source_path = Path("evaluation/test_contracts/software_service_contract.txt")
    review = ContractReviewService().review(source_path)
    client = FakeLLMClient(revised)

    result = AIContractRewriter(client).rewrite(original, review)
    text = document_text(result)

    assert "第一条 服务内容保持不变。" in text
    assert "【AI修改】第二条" in text
    assert "原合同全文" in client.messages[0].content


def test_ai_rewriter_rejects_unmarked_risky_contract() -> None:
    source_path = Path("evaluation/test_contracts/software_service_contract.txt")
    review = ContractReviewService().review(source_path)
    client = FakeLLMClient("软件服务合同\n" + "未标记的合同内容。" * 20)

    with pytest.raises(ContractRewriteError, match="未标记"):
        AIContractRewriter(client).rewrite("原合同" * 100, review)


def test_streamlit_generates_two_docx_downloads() -> None:
    source = Path("evaluation/test_contracts/software_service_contract.txt")
    original = source.read_text(encoding="utf-8")
    revised = original.replace("第一条", "【AI修改】第一条", 1)
    rewriter = AIContractRewriter(FakeLLMClient(revised))

    review, advice_document, revised_contract = generate_download_documents(
        source.name, source.read_bytes(), rewriter=rewriter
    )

    assert review.risks
    assert advice_document.startswith(b"PK")
    assert revised_contract.startswith(b"PK")
    assert "合同风险审查及修改建议书" in document_text(advice_document)
    assert "【AI修改】" in document_text(revised_contract)


def test_dashboard_metrics_use_existing_risk_scores() -> None:
    source = Path("evaluation/test_contracts/software_service_contract.txt")
    review = ContractReviewService().review(source)

    counts = risk_level_counts(review)

    assert contract_risk_score(review) == max(risk.risk_score for risk in review.risks)
    assert counts == {"HIGH": 8, "MIDDLE": 4, "LOW": 0}
    assert sum(counts.values()) == len(review.risks)


def test_analysis_process_reports_all_five_steps() -> None:
    source = Path("evaluation/test_contracts/software_service_contract.txt")
    original = source.read_text(encoding="utf-8")
    revised = original.replace("第一条", "【AI修改】第一条", 1)
    steps: list[tuple[int, str]] = []

    generate_download_documents(
        source.name,
        source.read_bytes(),
        rewriter=AIContractRewriter(FakeLLMClient(revised)),
        progress_callback=lambda step, message: steps.append((step, message)),
    )

    assert [step for step, _message in steps] == [1, 2, 3, 4, 5]
    assert steps[-1][1] == "生成报告完成"


def test_ai_revision_preview_counts_marked_paragraphs() -> None:
    content = AIContractRewriter._build_docx("合同\n【AI修改】第一条 修改一\n第二条 不变\n【AI修改】第三条 修改二")

    summary = ai_revision_summary(content)

    assert "2 处修改" in summary
    assert "第一条" in summary


def test_contract_diff_contains_original_revision_reason_and_reduction() -> None:
    source = Path("evaluation/test_contracts/software_service_contract.txt")
    original = source.read_text(encoding="utf-8")
    review = ContractReviewService().review(source)
    revised_text = original.replace("第一条", "【AI修改】第一条", 1)
    revised_contract = AIContractRewriter._build_docx(revised_text)

    content = ContractDiffGenerator().generate(original, revised_contract, review)
    text = document_text(content)

    assert content.startswith(b"PK")
    assert "合同修改说明" in text
    assert "原条款：" in text
    assert "修改后条款：" in text
    assert "修改理由：" in text
    assert "降低风险：" in text
    assert "CR-009" in text


@pytest.mark.parametrize(
    ("content", "message"),
    [
        (b"", "合同为空"),
        (b"x" * (MAX_UPLOAD_BYTES + 1), "超过 10 MB"),
    ],
    ids=["empty", "oversized"],
)
def test_upload_size_errors_are_chinese(content: bytes, message: str) -> None:
    with pytest.raises(UploadValidationError, match=message):
        validate_upload(content)
