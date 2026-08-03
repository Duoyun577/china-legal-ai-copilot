from pathlib import Path
from io import BytesIO
import sys

from docx import Document
import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "app"))

from streamlit_app import overall_risk_level, review_uploaded_contract
from document.contract_parser import ContractDocumentParser, DocumentParseError


def test_uploaded_contract_uses_existing_review_and_report_services() -> None:
    contract = Path("evaluation/test_contracts/software_service_contract.txt").read_bytes()

    review, report = review_uploaded_contract("software_service_contract.txt", contract)

    assert review.risks
    assert overall_risk_level(review) == "HIGH"
    assert "合同法律审查意见书" in report
    assert "CR-009" in report


def test_uploaded_filename_is_sanitized() -> None:
    review, report = review_uploaded_contract("../contract.txt", "普通服务合同".encode("utf-8"))

    assert review.contract_path.name == "contract.txt"
    assert report


def test_docx_contract_is_parsed_and_reviewed() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("软件开发服务合同")
    document.add_paragraph("付款期限由双方另行协商。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "验收"
    table.cell(0, 1).text = "测试合格后确认"
    document.save(buffer)

    review, report = review_uploaded_contract("contract.docx", buffer.getvalue())

    assert review.contract_type == "软件开发服务合同"
    assert any(risk.rule_id == "CR-009" for risk in review.risks)
    assert "contract.docx" in report


def test_pdf_parser_extracts_text_from_each_page(monkeypatch: pytest.MonkeyPatch) -> None:
    class Page:
        def __init__(self, text: str | None) -> None:
            self._text = text

        def extract_text(self) -> str | None:
            return self._text

    class Reader:
        def __init__(self, _stream: BytesIO) -> None:
            self.pages = [Page("软件服务合同"), Page(None), Page("付款期限不明确")]

    monkeypatch.setattr("document.contract_parser.PdfReader", Reader)

    parsed = ContractDocumentParser().parse("contract.PDF", b"fake pdf")

    assert parsed.file_type == "pdf"
    assert parsed.text == "软件服务合同\n\n付款期限不明确"


def test_image_only_pdf_has_clear_error(monkeypatch: pytest.MonkeyPatch) -> None:
    class Page:
        def extract_text(self) -> None:
            return None

    class Reader:
        def __init__(self, _stream: BytesIO) -> None:
            self.pages = [Page()]

    monkeypatch.setattr("document.contract_parser.PdfReader", Reader)

    with pytest.raises(DocumentParseError, match="扫描件"):
        ContractDocumentParser().parse("scan.pdf", b"fake pdf")


def test_unsupported_document_type_is_rejected() -> None:
    with pytest.raises(DocumentParseError, match="仅支持"):
        ContractDocumentParser().parse("contract.doc", b"content")
